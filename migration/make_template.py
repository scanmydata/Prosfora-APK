# -*- coding: utf-8 -*-
"""Παράγει το πρότυπο .docx της προσφοράς.

Το πρότυπο είναι *δεδομένα*, όχι κώδικας: ζει στο Drive του χρήστη και το
αλλάζει ελεύθερα. Αυτό εδώ φτιάχνει το εργοστασιακό — αυτό που ανεβαίνει την
πρώτη φορά και αυτό που επιστρέφει η «Επαναφορά εργοστασιακού προτύπου».

Σχεδιαστικοί στόχοι:
  * μία σελίδα A4 όσο το επιτρέπουν οι γραμμές της προσφοράς
  * παρατηρήσεις και τρόπος πληρωμής δίπλα-δίπλα, όχι σε στοίβα
  * το λογότυπο και τα χρώματά του δίνουν την ταυτότητα του εγγράφου

Τρέξιμο:  python migration/make_template.py
"""
import copy

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --- η παλέτα βγαίνει από το ίδιο το λογότυπο -------------------------------
# Το πράσινο και το μαύρο είναι αυτούσια από το tovapsimo-logo.png· το βαθύ
# πράσινο είναι η σκούρα εκδοχή του ίδιου τόνου, ώστε το κείμενο πάνω σε λευκό
# να διαβάζεται (το #00E2A2 είναι πολύ ανοιχτό για γράμματα).
GREEN = RGBColor(0x00, 0xE2, 0xA2)
DEEP = RGBColor(0x00, 0x6B, 0x4F)
INK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x59, 0x59, 0x59)

GREEN_HEX = "00E2A2"
DEEP_HEX = "006B4F"
TINT_HEX = "E9FCF5"       # πράσινο αραιωμένο, για γεμίσματα κελιών
LINE_HEX = "C9EFE2"       # διακριτικές γραμμές ανάμεσα στους χώρους

# Κρατιούνται τα παλιά ονόματα: το make_classic_template τα εισάγει από εδώ
NAVY = DEEP
ACCENT = GREEN

FONT = "Arial"          # ασφαλές για ελληνικά και στο Word και στο Google Docs

LOGO = "assets/branding/tovapsimo-logo.png"
OUT = "assets/pdf-template/ΠΡΟΣΦΟΡΑ ΕΛΑΙΟΧΡΩΜΑΤΙΣΜΩΝ.docx"

# Πόσος αέρας μπαίνει στην κορυφή των σελίδων μετά την πρώτη, σε στιγμές
TOP_AIR_PT = 34

# Τα τρία στοιχεία επικοινωνίας που κλείνουν και το κλασικό πρότυπο
LINKS = (
    ("📘", "www.facebook.com/tovapsimo"),
    ("🌐", "www.tovapsimo.gr"),
    ("📞", "6945773605"),
)


# --------------------------------------------------------------- helpers ---

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    previous = tcPr.find(qn("w:shd"))
    if previous is not None:                      # οι κλωνοποιημένες γραμμές φέρνουν ήδη γέμισμα
        tcPr.remove(previous)
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_color)
    tcPr.append(el)


def borders(cell, **edges):
    """edges: top/bottom/left/right -> (size_eighths, hex) ή None για κανένα."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    node = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        spec = edges.get(edge, "keep")
        if spec == "keep":
            continue
        e = OxmlElement("w:" + edge)
        if spec is None:
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
        else:
            size, color = spec
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(size))
            e.set(qn("w:color"), color)
        e.set(qn("w:space"), "0")
        node.append(e)
    tcPr.append(node)


def no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            borders(cell, top=None, left=None, bottom=None, right=None)


def cell_margins(table, top=40, bottom=40, left=80, right=80):
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for name, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement("w:" + name)
        e.set(qn("w:w"), str(value))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)


def rule(paragraph, hex_color, size=12):
    """Οριζόντια γραμμή ως κάτω περίγραμμα της παραγράφου."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    node = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    node.append(bottom)
    pPr.append(node)


def write(paragraph, text, size=9, bold=False, color=None, italic=False, spacing=None):
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if spacing is not None:                       # αραίωση γραμμάτων, σε twips
        rPr = run._r.get_or_add_rPr()
        el = OxmlElement("w:spacing")
        el.set(qn("w:val"), str(spacing))
        rPr.append(el)
    return run


def para(container, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p


def field(paragraph, instruction, placeholder, size, color):
    """Πεδίο του Word (PAGE, NUMPAGES) — υπολογίζεται από τον επεξεργαστή."""
    def run(child):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT)
        fonts.set(qn("w:hAnsi"), FONT)
        rPr.append(fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), color)
        rPr.append(col)
        r.append(rPr)
        r.append(child)
        paragraph._p.append(r)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run(sep)

    # Η τιμή που βλέπει όποιος δεν υπολογίζει το πεδίο· ο επεξεργαστής τη διορθώνει
    cached = OxmlElement("w:t")
    cached.text = placeholder
    run(cached)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run(end)


def page_furniture(section):
    """Αρίθμηση κάτω δεξιά, και αέρας στην κορυφή από τη 2η σελίδα και μετά.

    Το περιθώριο της σελίδας είναι ίδιο σε όλες τις σελίδες — δεν αλλάζει ανά
    σελίδα. Ο χώρος στην κορυφή των επόμενων σελίδων βγαίνει από την κεφαλίδα:
    με «διαφορετική πρώτη σελίδα», η πρώτη έχει κεφαλίδα μηδενικού ύψους ενώ οι
    υπόλοιπες μια κενή που σπρώχνει το κείμενο πιο κάτω. Αλλιώς το κείμενο της
    2ης σελίδας ξεκινά κολλητά στο πάνω χείλος.
    """
    section.different_first_page_header_footer = True

    first = section.first_page_header.paragraphs[0]
    first.paragraph_format.space_after = Pt(0)
    write(first, " ", size=1)

    rest = section.header.paragraphs[0]
    rest.paragraph_format.space_after = Pt(TOP_AIR_PT)
    write(rest, " ", size=1)

    for footer in (section.first_page_footer, section.footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        write(p, "Σελίδα ", size=7.5, color=GREY)
        field(p, " PAGE ", "1", size=7.5, color="595959")
        write(p, " από ", size=7.5, color=GREY)
        field(p, " NUMPAGES ", "1", size=7.5, color="595959")


def hanging(paragraph, indent_cm=0.45):
    """Κρεμαστή εσοχή: η δεύτερη γραμμή μιας παρατήρησης δεν πέφτει κάτω από την κουκκίδα."""
    pf = paragraph.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.first_line_indent = Cm(-indent_cm)


def vertical_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:vAlign")
    el.set(qn("w:val"), "center")
    tcPr.append(el)


def clone_row(table, row, before):
    """Αντίγραφο γραμμής πίνακα, τοποθετημένο πριν από τη [before].

    Το deepcopy κρατάει πλάτη, περιγράμματα και γεμίσματα — ξαναχτίζοντάς τα με
    το χέρι θα ξέφευγε κάποιο και η γραμμή θα έμοιαζε ξένη μέσα στον πίνακα.
    """
    new = copy.deepcopy(row._tr)
    before._tr.addprevious(new)
    from docx.table import _Row

    return _Row(new, table)


def set_cell_text(cell, text, bold=None, color=None, size=None):
    """Αλλάζει το κείμενο κρατώντας τη μορφοποίηση του πρώτου run."""
    p = cell.paragraphs[0]
    runs = p.runs
    if not runs:
        write(p, text, size=size or 9, bold=bool(bold), color=color)
        return
    runs[0].text = text
    for extra in runs[1:]:
        extra.text = ""
    if bold is not None:
        runs[0].bold = bold
    if color is not None:
        runs[0].font.color.rgb = color
    if size is not None:
        runs[0].font.size = Pt(size)


# ------------------------------------------------------------- το έγγραφο ---

def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(9)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    # Το πρότυπο του python-docx είναι Letter· χωρίς αυτό το PDF βγαίνει σε λάθος χαρτί
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)
    width = section.page_width - section.left_margin - section.right_margin

    page_furniture(section)

    # --- επικεφαλίδα: λογότυπο αριστερά, επικοινωνία δεξιά ----------------
    head = doc.add_table(rows=1, cols=2)
    head.alignment = WD_TABLE_ALIGNMENT.CENTER
    head.autofit = False
    no_borders(head)
    cell_margins(head, left=0, right=0)
    for index, w in ((0, Cm(6.4)), (1, Cm(11.4))):
        head.columns[index].width = w
        head.rows[0].cells[index].width = w

    left = head.rows[0].cells[0]
    left.paragraphs[0].text = ""
    vertical_center(left)
    left.paragraphs[0].add_run().add_picture(LOGO, width=Cm(4.4))

    right = head.rows[0].cells[1]
    right.paragraphs[0].text = ""
    vertical_center(right)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    write(right.paragraphs[0], "ΓΙΩΡΓΟΣ ΔΟΥΡΑΜΑΝΗΣ", size=13, bold=True, color=INK, spacing=20)
    p = para(right, align=WD_ALIGN_PARAGRAPH.RIGHT)
    write(p, "ΕΛΑΙΟΧΡΩΜΑΤΙΣΜΟΙ · ΑΝΑΚΑΙΝΙΣΕΙΣ · ΜΟΝΩΣΕΙΣ", size=7.5, color=GREY, spacing=24)
    p = para(right, align=WD_ALIGN_PARAGRAPH.RIGHT, before=4)
    write(p, "6945 773605 · tovapsimo.gr", size=8.5, bold=True, color=DEEP)

    p = para(doc, before=2, after=8)
    rule(p, GREEN_HEX, size=8)

    # --- τίτλος -----------------------------------------------------------
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    write(p, "ΠΡΟΣΦΟΡΑ ΕΛΑΙΟΧΡΩΜΑΤΙΣΜΩΝ", size=15, bold=True, color=DEEP, spacing=40)

    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    write(p, "<<[Είδος]>> ΕΠΙ ΤΗΣ ΟΔΟΥ <<[Οδός / Περιοχή]>>", size=11, bold=True, color=INK)

    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=9)
    write(p, "ΑΘΗΝΑ, <<[Ημερομηνία]>>", size=8.5, color=GREY)

    # --- ανάλυση χώρων ----------------------------------------------------
    cols = (Cm(9.0), Cm(2.6), Cm(2.9), Cm(3.3))
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell_margins(table, top=50, bottom=50, left=100, right=100)
    for index, w in enumerate(cols):
        table.columns[index].width = w
        for row in table.rows:
            row.cells[index].width = w

    headers = ("ΠΕΡΙΓΡΑΦΗ ΧΩΡΟΥ", "ΕΠΙΦΑΝΕΙΑ (Τ.Μ.)", "ΤΙΜΗ ΜΟΝΑΔΟΣ", "ΣΥΝΟΛΟ")
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, DEEP_HEX)
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        )
        write(cell.paragraphs[0], text, size=8, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        borders(cell, top=(6, DEEP_HEX), bottom=(6, DEEP_HEX), left=None, right=None)

    # η γραμμή-πρότυπο: επαναλαμβάνεται μία φορά ανά χώρο
    body = (
        "<<Start:[Χώροι]>><<[Περιγραφή Χώρου]>>",
        "<<[Επιφάνεια (τ.μ.)]>>",
        "<<[Τιμή Μονάδος]>>",
        "<<[Σύνολο Γραμμής]>><<End>>",
    )
    for i, text in enumerate(body):
        cell = table.rows[1].cells[i]
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        )
        write(cell.paragraphs[0], text, size=9, color=INK)
        borders(cell, top=None, bottom=(4, LINE_HEX), left=None, right=None)

    total_row = table.rows[2]
    for i in range(4):
        cell = total_row.cells[i]
        shade(cell, TINT_HEX)
        borders(cell, top=(10, DEEP_HEX), bottom=(10, DEEP_HEX), left=None, right=None)
        cell.paragraphs[0].text = ""
    total_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    write(total_row.cells[0].paragraphs[0], "ΓΕΝΙΚΟ ΣΥΝΟΛΟ", size=10, bold=True, color=DEEP)
    total_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    write(total_row.cells[3].paragraphs[0], "<<[Γενικό Σύνολο Live]>>", size=10,
          bold=True, color=DEEP)

    add_closing_rows(table, total_row)

    para(doc, after=8)

    # --- παρατηρήσεις + τρόπος πληρωμής, δίπλα-δίπλα ----------------------
    split = doc.add_table(rows=1, cols=2)
    split.alignment = WD_TABLE_ALIGNMENT.CENTER
    split.autofit = False
    no_borders(split)
    cell_margins(split, left=0, right=140)
    for index, w in ((0, Cm(11.4)), (1, Cm(6.4))):
        split.columns[index].width = w
        split.rows[0].cells[index].width = w

    notes = split.rows[0].cells[0]
    notes.paragraphs[0].text = ""
    write(notes.paragraphs[0], "ΠΑΡΑΤΗΡΗΣΕΙΣ", size=9, bold=True, color=DEEP, spacing=20)
    rule(notes.paragraphs[0], GREEN_HEX, size=6)
    p = para(notes, before=3)
    hanging(p)
    write(p, "•  <<[Παρατηρήσεις]>>", size=8.5, color=INK)

    pay = split.rows[0].cells[1]
    pay.paragraphs[0].text = ""
    write(pay.paragraphs[0], "ΤΡΟΠΟΣ ΠΛΗΡΩΜΗΣ", size=9, bold=True, color=DEEP, spacing=20)
    rule(pay.paragraphs[0], GREEN_HEX, size=6)
    p = para(pay, before=3)
    hanging(p)
    write(p, "•  <<[Τρόπος Πληρωμής]>>", size=8.5, color=INK)

    p = para(pay, before=7)
    write(p, "ΙΣΧΥΣ ΠΡΟΣΦΟΡΑΣ", size=9, bold=True, color=DEEP, spacing=20)
    rule(p, GREEN_HEX, size=6)
    p = para(pay, before=3)
    write(p, "Η προσφορά ισχύει έως <<[Ισχύει έως]>>", size=8.5, color=INK)

    para(doc, after=6)

    # --- η έμμεση διαφήμιση ----------------------------------------------
    strip = doc.add_table(rows=1, cols=1)
    strip.alignment = WD_TABLE_ALIGNMENT.CENTER
    strip.autofit = False
    strip.columns[0].width = width
    cell = strip.rows[0].cells[0]
    cell.width = width
    shade(cell, TINT_HEX)
    borders(cell, top=None, bottom=None, left=(30, GREEN_HEX), right=None)
    cell_margins(strip, top=70, bottom=70, left=160, right=120)
    cell.paragraphs[0].text = ""
    write(cell.paragraphs[0],
          "Φωτογραφίες από ολοκληρωμένα έργα, χρώματα και αξιολογήσεις πελατών:",
          size=8.5, color=GREY)
    # Τα ίδια τρία στοιχεία που κλείνουν και το κλασικό πρότυπο. Μπαίνουν μέσα
    # στη λωρίδα ώστε να διαβάζονται ως συνέχεια της πρότασης από πάνω.
    for icon, text in LINKS:
        p = para(cell, before=2)
        write(p, icon + "  ", size=8.5, color=INK)
        write(p, text, size=8.5, bold=True, color=DEEP)

    # --- υπογραφή ---------------------------------------------------------
    p = para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, before=14)
    write(p, "Ο ΕΡΓΟΛΗΠΤΗΣ", size=8.5, color=GREY, spacing=20)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, before=2)
    write(p, "ΓΙΩΡΓΟΣ ΔΟΥΡΑΜΑΝΗΣ", size=10, bold=True, color=INK)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT)
    write(p, "6945 773605", size=8.5, color=GREY)

    doc.save(OUT)
    return OUT


# Οι γραμμές που κλείνουν τον πίνακα, με τη σειρά που μπαίνουν πάνω από το
# γενικό σύνολο. Ο δείκτης «Αν …» λέει στην εφαρμογή πότε η γραμμή υπάρχει:
# ό,τι δεν ισχύει για τη συγκεκριμένη προσφορά σβήνεται ολόκληρο.
CLOSING_ROWS = (
    ("<<[Αν Σκαλωσιά]>>ΣΚΑΛΩΣΙΑ", "<<[Σκαλωσιά]>>", False),
    ("<<[Αν Άδεια]>>ΑΔΕΙΑ ΜΙΚΡΗΣ ΚΛΙΜΑΚΑΣ ΕΡΓΑΣΙΩΝ", "<<[Άδεια]>>", False),
    ("<<[Αν ΦΠΑ]>>ΣΥΝΟΛΟ", "<<[Καθαρή Αξία]>>", True),
    ("<<[Αν ΦΠΑ]>>ΦΠΑ 24%", "<<[ΦΠΑ]>>", False),
)


def add_closing_rows(table, total_row):
    """Πρόσθετα κόστη, σύνολο και ΦΠΑ, ακριβώς πάνω από το γενικό σύνολο.

    Το «ΣΥΝΟΛΟ» είναι έντονο όπως και το «ΓΕΝΙΚΟ ΣΥΝΟΛΟ»: είναι τα δύο ποσά που
    διαβάζει κανείς πρώτα — το ένα χωρίς ΦΠΑ και το άλλο με.
    """
    for label, value, bold in CLOSING_ROWS:
        row = clone_row(table, total_row, total_row)
        for i in range(4):
            cell = row.cells[i]
            shade(cell, "FFFFFF")
            borders(cell, top=None, bottom=(4, LINE_HEX), left=None, right=None)
            cell.paragraphs[0].text = ""
        colour = DEEP if bold else GREY
        size = 9.5 if bold else 8.5
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        write(row.cells[0].paragraphs[0], label, size=size, bold=bold, color=colour)
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        write(row.cells[3].paragraphs[0], value, size=size, bold=bold, color=colour)


if __name__ == "__main__":
    print("γράφτηκε:", build())
