# -*- coding: utf-8 -*-
"""Ετοιμάζει το «κλασικό» πρότυπο — αυτό με τα λογότυπα και τις φωτογραφίες.

Πηγή είναι το πρότυπο που ερχόταν από το AppSheet. Δεν ξαναφτιάχνεται από την
αρχή: κρατιέται όπως είναι (εικόνες, δείγματα εργασιών, στοιχεία επικοινωνίας)
και του προστίθενται μόνο όσα λείπουν —

  * αρίθμηση σελίδων κάτω δεξιά, σε κάθε σελίδα
  * αέρας στην κορυφή από τη 2η σελίδα και μετά
  * ισχύς προσφοράς και τρόπος πληρωμής
  * σκαλωσιά, άδεια μικρής κλίμακας, σύνολο και ΦΠΑ 24% μέσα στον πίνακα
  * στοιχίσεις που έβγαιναν στραβές στο τυπωμένο PDF

Τρέξιμο:
    python migration/make_classic_template.py --source "…/ΠΡΟΣΦΟΡΑ ΕΛΑΙΟΧΡΩΜΑΤΙΣΜΩΝ.docx"
"""
import argparse
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips
from docx.table import Table

from make_template import (
    CLOSING_ROWS, DEEP, FONT, GREEN_HEX, GREY, clone_row, field, para, rule,
    set_cell_text, write,
)

OUT = "assets/pdf-template/ΠΡΟΣΦΟΡΑ ΕΛΑΙΟΧΡΩΜΑΤΙΣΜΩΝ — κλασικό.docx"

# Πόσος αέρας στην κορυφή των σελίδων μετά την πρώτη, σε στιγμές
TOP_AIR_PT = 34

# Το μπλε των επικεφαλίδων του κλασικού. Ορίζεται εδώ και δεν κληρονομείται από
# το άλλο πρότυπο, που κινείται πια στα χρώματα του λογοτύπου.
HEADING = RGBColor(0x1F, 0x38, 0x64)


def furnish(section):
    """Αρίθμηση σελίδων και αέρας στις επόμενες σελίδες.

    Ίδια λύση με το άλλο πρότυπο: τα περιθώρια δεν αλλάζουν ανά σελίδα, οπότε ο
    χώρος στην κορυφή βγαίνει από κεφαλίδα που υπάρχει μόνο στις σελίδες μετά
    την πρώτη.
    """
    section.different_first_page_header_footer = True
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    first = section.first_page_header.paragraphs[0]
    first.paragraph_format.space_after = Pt(0)
    write(first, " ", size=1)

    rest = section.header.paragraphs[0]
    rest.paragraph_format.space_after = Pt(TOP_AIR_PT)
    write(rest, " ", size=1)

    for footer in (section.first_page_footer, section.footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        write(p, "Σελίδα ", size=8, color=GREY)
        field(p, " PAGE ", "1", size=8, color="595959")
        write(p, " από ", size=8, color=GREY)
        field(p, " NUMPAGES ", "1", size=8, color="595959")


def find_paragraph(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def insert_before(anchor, doc, builder):
    """Φτιάχνει παράγραφο στο τέλος και τη μετακινεί πριν το [anchor]."""
    p = builder(doc)
    anchor._p.addprevious(p._p)
    return p


def add_terms(doc):
    """Ισχύς και τρόπος πληρωμής, πριν από τα δείγματα εργασιών."""
    anchor = find_paragraph(doc, "ΔΕΙΓΜΑΤΑ ΕΡΓΑΣΙΩΝ")
    if anchor is None:
        raise SystemExit("δεν βρέθηκε το σημείο «ΔΕΙΓΜΑΤΑ ΕΡΓΑΣΙΩΝ»")

    def heading(text):
        def build(d):
            p = para(d, before=8, after=2)
            write(p, text, size=10, bold=True, color=HEADING, spacing=20)
            # Πράσινο του λογοτύπου, όχι το πορτοκαλί που ήταν πριν
            rule(p, GREEN_HEX, size=6)
            return p
        return build

    def line(text, indent=0.45):
        def build(d):
            p = para(d, after=1)
            if indent:
                p.paragraph_format.left_indent = Cm(indent)
                p.paragraph_format.first_line_indent = Cm(-indent)
            write(p, text, size=10)
            return p
        return build

    insert_before(anchor, doc, heading("ΤΡΟΠΟΣ ΠΛΗΡΩΜΗΣ"))
    insert_before(anchor, doc, line("•  <<[Τρόπος Πληρωμής]>>"))
    insert_before(anchor, doc, heading("ΙΣΧΥΣ ΠΡΟΣΦΟΡΑΣ"))
    insert_before(anchor, doc, line("Η προσφορά ισχύει έως <<[Ισχύει έως]>>", indent=0))
    insert_before(anchor, doc, lambda d: para(d, after=6))


def centre_everything(table):
    """Κεντράρει κάθε παράγραφο του πίνακα — εικόνες και λεζάντες μαζί."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def table_after(paragraph, doc):
    """Ο πρώτος πίνακας που ακολουθεί την [paragraph]."""
    node = paragraph._p.getnext()
    while node is not None and node.tag != qn("w:tbl"):
        node = node.getnext()
    return None if node is None else Table(node, doc)


def prices_table(doc):
    """Ο πίνακας της ανάλυσης χώρων — τον βρίσκουμε από τις επικεφαλίδες του."""
    for table in doc.tables:
        head = " ".join(cell.text.strip() for cell in table.rows[0].cells)
        if "ΠΕΡΙΓΡΑΦΗ ΧΩΡΟΥ" in head and "ΣΥΝΟΛΟ" in head:
            return table
    return None


def table_width(table):
    """Πλάτος του πίνακα σε twips, από το `w:tblW` ή από το άθροισμα του grid."""
    tblPr = table._tbl.tblPr
    node = tblPr.find(qn("w:tblW")) if tblPr is not None else None
    if node is not None and node.get(qn("w:type")) == "dxa":
        return int(float(node.get(qn("w:w"))))
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return None
    return int(sum(float(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol"))))


def align_date_with_table(doc, section):
    """Η ημερομηνία δεξιά, αλλά στο ίδιο δεξί χείλος με τον πίνακα.

    Σκέτο «δεξιά» θα την πήγαινε ως το περιθώριο της σελίδας, ενώ ο πίνακας
    σταματάει λίγο πριν — και η ημερομηνία θα κρεμόταν πιο έξω από τη στήλη
    «ΣΥΝΟΛΟ». Το κενό ανάμεσά τους μπαίνει ως δεξιά εσοχή της παραγράφου.
    """
    date = None
    for p in doc.paragraphs:
        if "<<[Ημερομηνία]>>" in p.text:
            date = p
            break
    if date is None:
        return

    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = prices_table(doc)
    width = table_width(table) if table is not None else None
    if width is None:
        return
    # Η αφαίρεση δύο Length δίνει σκέτο int σε EMU — 1 twip = 635 EMU
    content = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    gap = content // 635 - width
    if gap > 0:
        date.paragraph_format.right_indent = Twips(gap)


def add_closing_rows(doc):
    """Πρόσθετα κόστη, σύνολο και ΦΠΑ, ακριβώς πάνω από το γενικό σύνολο.

    Οι γραμμές γράφονται πάντα στο πρότυπο· η εφαρμογή αφαιρεί όποια δεν ισχύει
    για τη συγκεκριμένη προσφορά — αυτό δηλώνει ο δείκτης «Αν …» στην αρχή τους.
    """
    table = prices_table(doc)
    if table is None:
        raise SystemExit("δεν βρέθηκε ο πίνακας της ανάλυσης χώρων")

    total_row = table.rows[-1]
    if "<<[Αν " in total_row.cells[0].text:
        return                                     # έχει ξανατρέξει

    # Στο πρωτότυπο η τελευταία γραμμή λέει σκέτο «ΣΥΝΟΛΟ»· με τον ΦΠΑ από
    # πάνω της χρειάζεται να ξεχωρίζει ότι αυτή είναι η τελική
    set_cell_text(total_row.cells[0], "ΓΕΝΙΚΟ ΣΥΝΟΛΟ")

    for label, value, bold in CLOSING_ROWS:
        row = clone_row(table, total_row, total_row)
        colour = DEEP if bold else GREY
        set_cell_text(row.cells[0], label, bold=bold, color=colour, size=9.5 if bold else 9)
        set_cell_text(row.cells[-1], value, bold=bold, color=colour, size=9.5 if bold else 9)


def fix_alignment(doc):
    """Οι στοιχίσεις των δειγμάτων εργασιών.

    Οι εικόνες κάθονταν αριστερά μέσα στα κελιά τους, με μία λεζάντα να μη
    συμφωνεί με τις άλλες δύο.
    """
    samples = find_paragraph(doc, "ΔΕΙΓΜΑΤΑ ΕΡΓΑΣΙΩΝ")
    if samples is None:
        return
    samples.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = table_after(samples, doc)
    if table is None:
        return
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    centre_everything(table)


def build(source):
    doc = Document(source)

    section = doc.sections[0]
    # Το πρωτότυπο είναι ήδη A4, αλλά ας μην εξαρτιόμαστε από αυτό
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Πάνω/κάτω ήταν μισό εκατοστό — πολύ κολλητά για εκτύπωση
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    furnish(section)
    add_terms(doc)
    add_closing_rows(doc)
    align_date_with_table(doc, section)
    fix_alignment(doc)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    args = parser.parse_args()
    print("γράφτηκε:", build(args.source))
